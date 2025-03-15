import sys
import os
import json
import subprocess
from docx import Document
from docx2pdf import convert
import shutil

def convert_to_pdf_with_libreoffice(input_docx, output_pdf):
    libreoffice_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    temp_dir = os.path.dirname(output_pdf)  # Use same directory as output PDF
    
    # Convert DOCX to PDF in the same directory
    subprocess.run([libreoffice_path, "--headless", "--convert-to", "pdf", "--outdir", temp_dir, input_docx], check=True)

    # Rename the output file
    generated_pdf = os.path.join(temp_dir, os.path.splitext(os.path.basename(input_docx))[0] + ".pdf")
    if os.path.exists(generated_pdf):
        shutil.move(generated_pdf, output_pdf)
        print(f"PDF saved as {output_pdf}")

def generate_invoice(json_file): 
    # Check if JSON file path is provided as an argument

    if not os.path.exists(json_file):
        print(f"Error: JSON file '{json_file}' not found.")
        sys.exit(1)

    # Load the .docx template
    # Define directory dynamically
    script_dir = os.path.dirname(os.path.abspath(__file__))  # Get current script directory
    template_path = os.path.join(script_dir, "..", "invoice_generator", "template_invoice.docx")

    output_dir = os.path.join(script_dir, "..", "invoices/")
    os.makedirs(output_dir, exist_ok=True)
    doc = Document(template_path)

    # Load JSON data
    with open(json_file, "r") as f:
        invoice_data = json.load(f)

    # # Replace placeholders in paragraphs
    # for para in doc.paragraphs:
    #     text = para.text
    #     text = text.replace("Invoice #001", f"Invoice #{invoice_data['Invoice Number']}")
    #     text = text.replace("Date:12-30-23", f"Date: {invoice_data['Invoice Date']}")
    #     para.text = text  # Update the paragraph text

    # Replace placeholders inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace("Invoice #001", f"Invoice #{invoice_data['Invoice Number']}")
                cell.text = cell.text.replace("Date:12-30-23", f"Date: {invoice_data['Invoice Date']}")
                cell.text = cell.text.replace("Billing Address", invoice_data["Billing Address"])
                cell.text = cell.text.replace("Shipping Address", invoice_data["Shipping Address"])
                cell.text = cell.text.replace("Instructions", invoice_data["Instructions"])
                cell.text = cell.text.replace("Total Amount", invoice_data.get("Total Amount").split(":")[1])

    # Find the invoice items table
    items_table = None
    for table in doc.tables:
        for row in table.rows:
            if "Quantity" in row.cells[0].text:  # Identify the correct table by header
                items_table = table  # Save reference to the correct table
                break
        if items_table:
            break  # Stop searching if we found the table

    if items_table:
        item_rows = items_table.rows[1:]  # Skip the header row
        item_index = 0  # Keep track of which item we're filling

        # Fill pre-existing empty rows first
        for row in item_rows:
            if item_index < len(invoice_data["Items"]):
                item = invoice_data["Items"][item_index]
                row.cells[0].text = item["Quantity"]
                row.cells[1].text = item["Description"]
                row.cells[2].text = "$" + item["Unit Price"]
                row.cells[3].text = "$" + item["Total"]
                item_index += 1
            else:
                break  # Stop if all items are added

        # If there are more items than available rows, add new rows
        while item_index < len(invoice_data["Items"]):
            item = invoice_data["Items"][item_index]
            new_row = items_table.add_row().cells  # Add new row
            new_row[0].text = item["Quantity"]
            new_row[1].text = item["Description"]
            new_row[2].text = item["Unit Price"]
            new_row[3].text = "$"  + item["Total"]
            item_index += 1

    # Save the filled document
    output_docx = os.path.join(output_dir, f"invoice_{invoice_data['Invoice Number']}.docx")
    doc.save(output_docx)
    print(f"Invoice saved as {output_docx}")

    # Convert to PDF
    output_pdf = os.path.join(output_dir, f"invoice_{invoice_data['Invoice Number']}.pdf")
    try: 
        print(f"Running it on libreoffice")
        convert_to_pdf_with_libreoffice(output_docx, output_pdf)  
    except: 
        print(f"Running it on docx2pdf")
        convert(output_docx, output_pdf)
        print(f"PDF saved as {output_pdf}")
        

# Generate the invoice 
if __name__ == "__main__":
    generate_invoice()

